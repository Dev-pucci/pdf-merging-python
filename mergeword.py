import os
from docx import Document
from fpdf import FPDF

def load_document(file_path):
    """Load a Word document and handle errors."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file '{file_path}' does not exist. Check the path.")
    if not file_path.endswith(".docx"):
        raise ValueError(f"The file '{file_path}' is not a .docx file. Convert it to .docx format.")
    print(f"Loading file: {file_path}")  # Debugging
    return Document(file_path)

def merge_word_documents(doc1_path, doc2_path, output_docx_path):
    """Merge two Word documents into one."""
    # Load the Word documents
    doc1 = load_document(doc1_path)
    doc2 = load_document(doc2_path)

    # Append contents of doc2 to doc1
    for element in doc2.element.body:
        doc1.element.body.append(element)

    # Save the merged Word document
    doc1.save(output_docx_path)
    print(f"Merged Word document saved as: {output_docx_path}")

def create_pdf_from_docx(input_docx_path, output_pdf_path):
    """Create a PDF from a Word document."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Open the merged Word document and add its paragraphs to the PDF
    merged_doc = Document(input_docx_path)
    for para in merged_doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            pdf.multi_cell(0, 10, para.text)
            pdf.ln(5)  # Add space between paragraphs

    # Save the PDF
    pdf.output(output_pdf_path)
    print(f"Merged PDF saved as: {output_pdf_path}")

# Input file paths
doc1_path = r"c:\Users\dapu\Downloads\IFATAfrica25 (2).docx- FILLED"
doc2_path = r"c:\Users\dapu\Downloads\OP - 2024 (1).docx"

# Output file paths
output_docx_path = r"c:\Users\dapu\Downloads\Merged_Document.docx"
output_pdf_path = r"c:\Users\dapu\Downloads\Merged_Document.pdf"

try:
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)

    # Run the merging process
    merge_word_documents(doc1_path, doc2_path, output_docx_path)
    create_pdf_from_docx(output_docx_path, output_pdf_path)
except (FileNotFoundError, ValueError) as e:
    print(f"Error: {e}")
except Exception as e:
    print(f"An unexpected error occurred: {e}")